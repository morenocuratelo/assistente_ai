"""
Archive service implementation.
Handles document processing, organization, and advanced search capabilities.
"""

import os
import hashlib
import mimetypes
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import asyncio
import logging

from ...database.repositories.document_repository import DocumentRepository
from ...database.models.base import Document, ProcessingStatus, DocumentResponse
from ...core.errors.error_handler import handle_errors, handle_errors_async
from ...core.errors.error_types import (
    DocumentProcessingError,
    FileNotFoundError,
    ValidationError,
    TextExtractionError
)


class ArchiveService:
    """Service for document archive management."""

    def __init__(self, db_path: str, upload_dir: str = "documenti_da_processare"):
        """Initialize archive service.

        Args:
            db_path: Path to database file
            upload_dir: Directory for uploaded files
        """
        self.db_path = db_path
        self.upload_dir = Path(upload_dir)
        self.document_repository = DocumentRepository(db_path)
        self.logger = logging.getLogger(__name__)

        # Ensure upload directory exists
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    @handle_errors(operation="process_document", component="archive_service")
    def process_document(
        self,
        file_path: str,
        project_id: str,
        user_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> DocumentResponse:
        """Process a single document.

        Args:
            file_path: Path to document file
            project_id: Project ID for organization
            user_id: User ID who uploaded the document
            metadata: Additional metadata

        Returns:
            Document response with processing results

        Raises:
            DocumentProcessingError: If processing fails
        """
        file_path = Path(file_path)

        # Validate file exists
        if not file_path.exists():
            raise FileNotFoundError(str(file_path))

        # Validate file size
        file_size = file_path.stat().st_size
        max_size = 100 * 1024 * 1024  # 100MB
        if file_size > max_size:
            raise ValidationError(
                f"File too large: {file_size} bytes",
                "file_size",
                file_size
            )

        # Extract file information
        file_name = file_path.name
        mime_type = mimetypes.guess_type(file_name)[0] or "application/octet-stream"
        content_hash = self._calculate_content_hash(file_path)

        # Check for duplicates
        existing_docs = self.document_repository.get_by_content_hash(content_hash)
        if existing_docs:
            self.logger.info(f"Duplicate document found: {file_name}")
            # Return existing document info instead of processing again
            existing = existing_docs[0]
            return DocumentResponse(
                document=existing,
                processing_time_ms=0,
                word_count=0,
                ai_confidence=0.0
            )

        # Create document record
        document = Document(
            project_id=project_id,
            file_name=file_name,
            file_size=file_size,
            mime_type=mime_type,
            content_hash=content_hash,
            created_by=int(user_id) if user_id else None,
            **(metadata or {})
        )

        # Save to database
        saved_document = self.document_repository.create(document)

        # Extract text and generate preview
        try:
            text_content = self._extract_text_content(file_path)
            preview = self._generate_preview(text_content)
            keywords = self._extract_keywords(text_content)

            # Update document with extracted information
            update_data = {
                'formatted_preview': preview,
                'keywords': keywords,
                'processing_status': ProcessingStatus.COMPLETED,
                'ai_tasks': {
                    'text_extraction': True,
                    'preview_generation': True,
                    'keyword_extraction': True,
                    'processing_timestamp': datetime.utcnow().isoformat()
                }
            }

            temp_doc = Document(**update_data)
            self.document_repository.update(saved_document.id, temp_doc)

            # Reload updated document
            saved_document = self.document_repository.get_by_id(saved_document.id)

        except Exception as e:
            # Mark as failed but keep the record
            error_msg = f"Processing failed: {str(e)}"
            self.document_repository.update_processing_status(
                saved_document.id,
                ProcessingStatus.FAILED,
                error_msg
            )
            self.logger.error(f"Document processing failed for {file_name}: {e}")
            raise DocumentProcessingError(error_msg, str(file_path))

        # Create response
        return DocumentResponse(
            document=saved_document,
            processing_time_ms=0,  # Could be measured
            word_count=len(text_content.split()) if 'text_content' in locals() else 0,
            ai_confidence=0.8  # Could be calculated based on extraction quality
        )

    @handle_errors(operation="batch_process_documents", component="archive_service")
    def batch_process_documents(
        self,
        file_paths: List[str],
        project_id: str,
        user_id: Optional[str] = None
    ) -> List[DocumentResponse]:
        """Process multiple documents in batch.

        Args:
            file_paths: List of file paths to process
            project_id: Project ID for organization
            user_id: User ID who uploaded the documents

        Returns:
            List of document responses
        """
        results = []

        for file_path in file_paths:
            try:
                result = self.process_document(file_path, project_id, user_id)
                results.append(result)
            except Exception as e:
                self.logger.error(f"Batch processing failed for {file_path}: {e}")
                # Create failed document record for tracking
                failed_doc = Document(
                    project_id=project_id,
                    file_name=Path(file_path).name,
                    processing_status=ProcessingStatus.FAILED,
                    created_by=int(user_id) if user_id else None
                )
                self.document_repository.create(failed_doc)

                # Add failed result
                results.append(DocumentResponse(
                    document=failed_doc,
                    processing_time_ms=0,
                    word_count=0,
                    ai_confidence=0.0
                ))

        return results

    @handle_errors(operation="search_documents", component="archive_service")
    def search_documents(
        self,
        query: str,
        project_id: str,
        filters: Optional[Dict[str, Any]] = None,
        limit: int = 50
    ) -> List[DocumentResponse]:
        """Search documents with advanced filtering.

        Args:
            query: Search query text
            project_id: Project ID to search in
            filters: Additional filters
            limit: Maximum results

        Returns:
            List of matching documents
        """
        # Use repository advanced search
        search_results = self.document_repository.get_advanced_search(
            query=query,
            project_id=project_id,
            limit=limit,
            **(filters or {})
        )

        # Convert to response format
        responses = []
        for doc in search_results['documents']:
            responses.append(DocumentResponse(
                document=doc,
                processing_time_ms=search_results['query_time_ms'],
                word_count=self._estimate_word_count(doc.formatted_preview or ""),
                ai_confidence=0.8
            ))

        return responses

    @handle_errors(operation="organize_documents", component="archive_service")
    def organize_documents(
        self,
        project_id: str,
        criteria: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Organize documents based on criteria.

        Args:
            project_id: Project ID
            criteria: Organization criteria

        Returns:
            Organization results
        """
        # Get documents to organize
        documents = self.document_repository.get_by_project(project_id)

        organization_results = {
            'total_processed': 0,
            'by_category': {},
            'by_size': {},
            'by_date': {},
            'errors': []
        }

        for doc in documents:
            try:
                # Apply organization logic based on criteria
                category = self._determine_category(doc, criteria)
                size_bucket = self._determine_size_bucket(doc.file_size or 0)
                date_bucket = self._determine_date_bucket(doc.created_at)

                # Update organization results
                organization_results['by_category'][category] = \
                    organization_results['by_category'].get(category, 0) + 1
                organization_results['by_size'][size_bucket] = \
                    organization_results['by_size'].get(size_bucket, 0) + 1
                organization_results['by_date'][date_bucket] = \
                    organization_results['by_date'].get(date_bucket, 0) + 1

                organization_results['total_processed'] += 1

            except Exception as e:
                organization_results['errors'].append({
                    'document_id': doc.id,
                    'file_name': doc.file_name,
                    'error': str(e)
                })

        return organization_results

    @handle_errors(operation="batch_operation", component="archive_service")
    def batch_operation(
        self,
        operation: str,
        document_ids: List[int],
        project_id: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Execute batch operation on documents.

        Args:
            operation: Operation to perform (delete, move, tag, etc.)
            document_ids: List of document IDs
            project_id: Project ID
            **kwargs: Operation-specific parameters

        Returns:
            Batch operation results
        """
        results = {
            'operation': operation,
            'total_requested': len(document_ids),
            'successful': 0,
            'failed': 0,
            'errors': []
        }

        for doc_id in document_ids:
            try:
                if operation == "delete":
                    success = self.document_repository.delete(doc_id)
                elif operation == "archive":
                    # Mark as archived
                    temp_doc = Document(processing_status=ProcessingStatus.COMPLETED)
                    success = self.document_repository.update(doc_id, temp_doc)
                elif operation == "reprocess":
                    # Reset for reprocessing
                    temp_doc = Document(processing_status=ProcessingStatus.PENDING)
                    success = self.document_repository.update(doc_id, temp_doc)
                else:
                    results['errors'].append({
                        'document_id': doc_id,
                        'error': f"Unsupported operation: {operation}"
                    })
                    results['failed'] += 1
                    continue

                if success:
                    results['successful'] += 1
                else:
                    results['failed'] += 1

            except Exception as e:
                results['errors'].append({
                    'document_id': doc_id,
                    'error': str(e)
                })
                results['failed'] += 1

        return results

    def _calculate_content_hash(self, file_path: Path) -> str:
        """Calculate content hash for duplicate detection."""
        try:
            hash_obj = hashlib.sha256()
            with open(file_path, 'rb') as f:
                # Read in chunks for large files
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_obj.update(chunk)
            return hash_obj.hexdigest()
        except Exception as e:
            self.logger.error(f"Error calculating hash for {file_path}: {e}")
            return ""

    def _extract_text_content(self, file_path: Path) -> str:
        """Extract text content from file."""
        try:
            file_extension = file_path.suffix.lower()

            if file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_text_from_docx(file_path)
            else:
                # For unsupported formats, return filename as content
                return f"Content extraction not supported for {file_extension}"

        except Exception as e:
            raise TextExtractionError(str(file_path), "auto_detection")

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2

            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"

            return text.strip()

        except ImportError:
            self.logger.warning("PyPDF2 not available for PDF text extraction")
            return f"PDF file: {file_path.name}"
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return f"PDF file: {file_path.name}"

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text.strip()

        except ImportError:
            self.logger.warning("python-docx not available for DOCX text extraction")
            return f"Word document: {file_path.name}"
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return f"Word document: {file_path.name}"

    def _generate_preview(self, text_content: str, max_length: int = 500) -> str:
        """Generate preview text from content."""
        if not text_content:
            return ""

        # Clean and truncate text
        preview = text_content.strip()
        preview = ' '.join(preview.split())  # Remove extra whitespace

        if len(preview) <= max_length:
            return preview

        # Truncate at word boundary
        truncated = preview[:max_length]
        last_space = truncated.rfind(' ')

        if last_space > max_length * 0.8:
            return truncated[:last_space] + "..."
        else:
            return truncated + "..."

    def _extract_keywords(self, text_content: str, max_keywords: int = 10) -> List[str]:
        """Extract keywords from text content."""
        if not text_content:
            return []

        try:
            # Simple keyword extraction based on word frequency
            import re
            from collections import Counter

            # Remove punctuation and convert to lowercase
            words = re.findall(r'\b\w+\b', text_content.lower())

            # Filter out common stop words
            stop_words = {
                'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have',
                'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
                'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those'
            }

            filtered_words = [word for word in words if len(word) > 3 and word not in stop_words]

            # Count frequency
            word_counts = Counter(filtered_words)

            # Return top keywords
            return [word for word, count in word_counts.most_common(max_keywords)]

        except Exception as e:
            self.logger.error(f"Error extracting keywords: {e}")
            return []

    def _determine_category(self, document: Document, criteria: Dict[str, Any]) -> str:
        """Determine document category based on criteria."""
        # Simple category determination based on file extension and content
        file_extension = Path(document.file_name).suffix.lower()

        if file_extension == '.pdf':
            return 'academic_papers'
        elif file_extension in ['.docx', '.doc']:
            return 'documents'
        elif file_extension == '.txt':
            return 'notes'
        else:
            return 'other'

    def _determine_size_bucket(self, file_size: int) -> str:
        """Determine size bucket for document."""
        if file_size < 1024 * 1024:  # < 1MB
            return 'small'
        elif file_size < 10 * 1024 * 1024:  # < 10MB
            return 'medium'
        else:
            return 'large'

    def _determine_date_bucket(self, created_at: datetime) -> str:
        """Determine date bucket for document."""
        if isinstance(created_at, str):
            created_at = datetime.fromisoformat(created_at)

        now = datetime.utcnow()
        days_diff = (now - created_at).days

        if days_diff <= 7:
            return 'week'
        elif days_diff <= 30:
            return 'month'
        elif days_diff <= 365:
            return 'year'
        else:
            return 'older'

    def _estimate_word_count(self, text: str) -> int:
        """Estimate word count from text."""
        if not text:
            return 0
        return len(text.split())

    def get_archive_statistics(self, project_id: str) -> Dict[str, Any]:
        """Get comprehensive archive statistics.

        Args:
            project_id: Project ID

        Returns:
            Dictionary with archive statistics
        """
        try:
            # Get basic stats from repository
            basic_stats = self.document_repository.get_processing_stats(project_id)

            # Get storage summary
            storage_summary = self.document_repository.get_storage_summary(project_id)

            # Get recent activity
            recent_docs = self.document_repository.get_recent_documents(project_id, limit=10)

            return {
                'processing_stats': basic_stats,
                'storage_summary': storage_summary,
                'recent_documents_count': len(recent_docs),
                'last_updated': datetime.utcnow().isoformat(),
                'project_id': project_id
            }

        except Exception as e:
            self.logger.error(f"Error getting archive statistics: {e}")
            return {}

    def cleanup_old_documents(self, project_id: str, days_old: int = 365) -> int:
        """Clean up old documents.

        Args:
            project_id: Project ID
            days_old: Days after which documents are considered old

        Returns:
            Number of documents cleaned up
        """
        try:
            return self.document_repository.archive_old_documents(days_old, project_id)
        except Exception as e:
            self.logger.error(f"Error cleaning up old documents: {e}")
            return 0

    def find_duplicate_documents(self, project_id: str) -> List[Dict[str, Any]]:
        """Find duplicate documents in archive.

        Args:
            project_id: Project ID

        Returns:
            List of duplicate document groups
        """
        try:
            return self.document_repository.find_duplicates(project_id)
        except Exception as e:
            self.logger.error(f"Error finding duplicates: {e}")
            return []

    def get_documents_needing_reprocessing(self, project_id: str) -> List[Document]:
        """Get documents that need reprocessing.

        Args:
            project_id: Project ID

        Returns:
            List of documents needing reprocessing
        """
        try:
            # Get failed and pending documents
            failed_docs = self.document_repository.get_by_processing_status(
                ProcessingStatus.FAILED
            )
            pending_docs = self.document_repository.get_by_processing_status(
                ProcessingStatus.PENDING
            )

            # Filter by project
            project_docs = []
            for doc in failed_docs + pending_docs:
                if doc.project_id == project_id:
                    project_docs.append(doc)

            return project_docs

        except Exception as e:
            self.logger.error(f"Error getting documents needing reprocessing: {e}")
            return []

    def export_documents(
        self,
        document_ids: List[int],
        export_format: str = "json",
        include_content: bool = False
    ) -> str:
        """Export documents in specified format.

        Args:
            document_ids: List of document IDs to export
            export_format: Export format (json, csv, txt)
            include_content: Whether to include full content

        Returns:
            Export file path
        """
        try:
            # Get documents
            documents = []
            for doc_id in document_ids:
                doc = self.document_repository.get_by_id(doc_id)
                if doc:
                    documents.append(doc)

            if not documents:
                raise ValidationError("No documents found for export", "document_ids", document_ids)

            # Create export directory
            export_dir = Path("exports")
            export_dir.mkdir(exist_ok=True)

            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            filename = f"archive_export_{timestamp}.{export_format}"
            export_path = export_dir / filename

            if export_format.lower() == "json":
                self._export_as_json(documents, export_path, include_content)
            elif export_format.lower() == "csv":
                self._export_as_csv(documents, export_path, include_content)
            elif export_format.lower() == "txt":
                self._export_as_txt(documents, export_path, include_content)
            else:
                raise ValidationError(f"Unsupported export format: {export_format}", "export_format", export_format)

            self.logger.info(f"Exported {len(documents)} documents to {export_path}")
            return str(export_path)

        except Exception as e:
            self.logger.error(f"Error exporting documents: {e}")
            raise

    def _export_as_json(self, documents: List[Document], export_path: Path, include_content: bool) -> None:
        """Export documents as JSON."""
        export_data = []

        for doc in documents:
            doc_data = {
                'id': doc.id,
                'file_name': doc.file_name,
                'title': doc.title,
                'project_id': doc.project_id,
                'file_size': doc.file_size,
                'mime_type': doc.mime_type,
                'processing_status': doc.processing_status.value,
                'keywords': doc.keywords,
                'created_at': doc.created_at.isoformat() if doc.created_at else None,
                'updated_at': doc.updated_at.isoformat() if doc.updated_at else None,
            }

            if include_content and doc.formatted_preview:
                doc_data['content'] = doc.formatted_preview

            export_data.append(doc_data)

        with open(export_path, 'w', encoding='utf-8') as f:
            import json
            json.dump(export_data, f, indent=2, ensure_ascii=False)

    def _export_as_csv(self, documents: List[Document], export_path: Path, include_content: bool) -> None:
        """Export documents as CSV."""
        import csv

        with open(export_path, 'w', newline='', encoding='utf-8') as f:
            fieldnames = [
                'id', 'file_name', 'title', 'project_id', 'file_size',
                'mime_type', 'processing_status', 'keywords', 'created_at'
            ]

            if include_content:
                fieldnames.append('content')

            writer = csv.DictWriter(f, fieldnames=fieldnames)

            writer.writeheader()

            for doc in documents:
                row = {
                    'id': doc.id,
                    'file_name': doc.file_name,
                    'title': doc.title,
                    'project_id': doc.project_id,
                    'file_size': doc.file_size,
                    'mime_type': doc.mime_type,
                    'processing_status': doc.processing_status.value,
                    'keywords': ', '.join(doc.keywords),
                    'created_at': doc.created_at.isoformat() if doc.created_at else None,
                }

                if include_content and doc.formatted_preview:
                    row['content'] = doc.formatted_preview

                writer.writerow(row)

    def _export_as_txt(self, documents: List[Document], export_path: Path, include_content: bool) -> None:
        """Export documents as plain text."""
        with open(export_path, 'w', encoding='utf-8') as f:
            for doc in documents:
                f.write(f"=== Document: {doc.file_name} ===\n")
                f.write(f"Title: {doc.title or 'N/A'}\n")
                f.write(f"Size: {doc.file_size or 0} bytes\n")
                f.write(f"Status: {doc.processing_status.value}\n")
                f.write(f"Keywords: {', '.join(doc.keywords)}\n")

                if include_content and doc.formatted_preview:
                    f.write(f"\nContent:\n{doc.formatted_preview}\n")

                f.write("\n" + "="*50 + "\n\n")

    async def process_documents_async(
        self,
        file_paths: List[str],
        project_id: str,
        user_id: Optional[str] = None
    ) -> List[DocumentResponse]:
        """Process documents asynchronously.

        Args:
            file_paths: List of file paths to process
            project_id: Project ID for organization
            user_id: User ID who uploaded the documents

        Returns:
            List of document responses
        """
        # Split into batches for parallel processing
        batch_size = 5
        batches = [file_paths[i:i + batch_size] for i in range(0, len(file_paths), batch_size)]

        results = []

        for batch in batches:
            # Process batch in parallel
            batch_tasks = [
                self._process_single_async(file_path, project_id, user_id)
                for file_path in batch
            ]

            batch_results = await asyncio.gather(*batch_tasks, return_exceptions=True)

            for result in batch_results:
                if isinstance(result, Exception):
                    self.logger.error(f"Async processing failed: {result}")
                    # Create failed document record
                    failed_doc = Document(
                        project_id=project_id,
                        file_name=Path(batch[0]).name,  # Simplified
                        processing_status=ProcessingStatus.FAILED
                    )
                    results.append(DocumentResponse(
                        document=failed_doc,
                        processing_time_ms=0,
                        word_count=0,
                        ai_confidence=0.0
                    ))
                else:
                    results.append(result)

        return results

    async def _process_single_async(
        self,
        file_path: str,
        project_id: str,
        user_id: Optional[str]
    ) -> DocumentResponse:
        """Process single document asynchronously."""
        # Run sync processing in thread pool
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self.process_document,
            file_path,
            project_id,
            user_id
        )
