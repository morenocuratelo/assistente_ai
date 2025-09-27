"""
Task specifiche per Archivista AI
Implementazione delle task Celery per l'archiviazione e processamento documenti
"""

import os
import logging
from datetime import datetime
from typing import Dict, List, Optional, Any
import fitz  # PyMuPDF

# Configurazione logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def process_single_document(file_path: str, metadata: Optional[Dict] = None) -> Dict[str, Any]:
    """
    Processa un singolo documento

    Args:
        file_path: Percorso del file da processare
        metadata: Metadati aggiuntivi per il documento

    Returns:
        Dict con i risultati del processamento
    """
    try:
        logger.info(f"Inizio processamento documento: {file_path}")

        # Verifica che il file esista
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File non trovato: {file_path}")

        # Determina il tipo di file
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.pdf':
            result = process_pdf_document(file_path, metadata)
        elif file_extension in ['.docx', '.doc']:
            result = process_docx_document(file_path, metadata)
        else:
            raise ValueError(f"Tipo di file non supportato: {file_extension}")

        logger.info(f"Processamento completato per: {file_path}")
        return result

    except Exception as e:
        logger.error(f"Errore nel processamento del documento {file_path}: {str(e)}")
        return {
            'status': 'error',
            'file_path': file_path,
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }

def process_pdf_document(file_path: str, metadata: Optional[Dict] = None) -> Dict[str, Any]:
    """
    Processa un documento PDF

    Args:
        file_path: Percorso del file PDF
        metadata: Metadati del documento

    Returns:
        Dict con i risultati del processamento
    """
    try:
        from llama_parse import LlamaParse
        from llama_index.core import VectorStoreIndex, Document
        import fitz  # PyMuPDF

        # Estrai testo dal PDF
        text_content = extract_text_from_pdf(file_path)

        # Crea documento LlamaIndex
        document = Document(
            text=text_content,
            metadata={
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'file_type': 'pdf',
                'processed_at': datetime.now().isoformat(),
                **(metadata or {})
            }
        )

        # Crea indice
        index = VectorStoreIndex.from_documents([document])

        # Salva l'indice
        index.storage_context.persist(persist_dir="db_memoria")

        return {
            'status': 'success',
            'file_path': file_path,
            'file_type': 'pdf',
            'text_length': len(text_content),
            'document_count': 1,
            'timestamp': datetime.now().isoformat()
        }

    except Exception as e:
        logger.error(f"Errore nel processamento PDF {file_path}: {str(e)}")
        raise

def process_docx_document(file_path: str, metadata: Optional[Dict] = None) -> Dict[str, Any]:
    """
    Processa un documento DOCX

    Args:
        file_path: Percorso del file DOCX
        metadata: Metadati del documento

    Returns:
        Dict con i risultati del processamento
    """
    try:
        from docx import Document
        from llama_index.core import VectorStoreIndex, Document as LlamaDocument

        # Estrai testo dal DOCX
        doc = Document(file_path)
        text_content = extract_text_from_docx(doc)

        # Crea documento LlamaIndex
        document = LlamaDocument(
            text=text_content,
            metadata={
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'file_type': 'docx',
                'processed_at': datetime.now().isoformat(),
                **(metadata or {})
            }
        )

        # Crea indice
        index = VectorStoreIndex.from_documents([document])

        # Salva l'indice
        index.storage_context.persist(persist_dir="db_memoria")

        return {
            'status': 'success',
            'file_path': file_path,
            'file_type': 'docx',
            'text_length': len(text_content),
            'document_count': 1,
            'timestamp': datetime.now().isoformat()
        }

    except Exception as e:
        logger.error(f"Errore nel processamento DOCX {file_path}: {str(e)}")
        raise

def extract_text_from_pdf(file_path: str) -> str:
    """
    Estrae il testo da un file PDF

    Args:
        file_path: Percorso del file PDF

    Returns:
        str: Testo estratto dal PDF
    """
    try:
        doc = fitz.open(file_path)
        text = ""

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()

        doc.close()
        return text

    except Exception as e:
        logger.error(f"Errore nell'estrazione testo da PDF {file_path}: {str(e)}")
        raise

def extract_text_from_docx(doc) -> str:
    """
    Estrae il testo da un documento DOCX

    Args:
        doc: Oggetto Document di python-docx

    Returns:
        str: Testo estratto dal documento
    """
    try:
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        return text

    except Exception as e:
        logger.error(f"Errore nell'estrazione testo da DOCX: {str(e)}")
        raise

def batch_process_documents(file_paths: List[str], metadata_list: Optional[List[Dict]] = None) -> Dict[str, Any]:
    """
    Processa un batch di documenti

    Args:
        file_paths: Lista dei percorsi dei file da processare
        metadata_list: Lista opzionale di metadati per ogni file

    Returns:
        Dict con i risultati del batch processing
    """
    results = []
    successful = 0
    failed = 0

    for i, file_path in enumerate(file_paths):
        try:
            # Usa metadati se disponibili
            metadata = metadata_list[i] if metadata_list and i < len(metadata_list) else None

            result = process_single_document(file_path, metadata)
            results.append(result)

            if result['status'] == 'success':
                successful += 1
            else:
                failed += 1

        except Exception as e:
            failed += 1
            results.append({
                'status': 'error',
                'file_path': file_path,
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            })

    return {
        'status': 'completed',
        'total_files': len(file_paths),
        'successful': successful,
        'failed': failed,
        'results': results,
        'timestamp': datetime.now().isoformat()
    }

def generate_document_summary(document_id: str, max_length: int = 500) -> Dict[str, Any]:
    """
    Genera un riassunto di un documento

    Args:
        document_id: ID del documento
        max_length: Lunghezza massima del riassunto

    Returns:
        Dict con il riassunto generato
    """
    try:
        from llama_index.core import StorageContext, load_index_from_storage

        # Carica l'indice del documento
        storage_context = StorageContext.from_defaults(persist_dir="db_memoria")
        index = load_index_from_storage(storage_context)

        # Crea query engine per il riassunto
        query_engine = index.as_query_engine(
            response_mode="tree_summarize",
            similarity_top_k=3
        )

        # Genera il riassunto
        query = f"Genera un riassunto conciso del documento, lunghezza massima {max_length} caratteri."
        response = query_engine.query(query)

        return {
            'status': 'success',
            'document_id': document_id,
            'summary': str(response),
            'max_length': max_length,
            'timestamp': datetime.now().isoformat()
        }

    except Exception as e:
        logger.error(f"Errore nella generazione del riassunto per {document_id}: {str(e)}")
        return {
            'status': 'error',
            'document_id': document_id,
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }

def search_similar_documents(query: str, top_k: int = 5) -> Dict[str, Any]:
    """
    Cerca documenti simili ad una query

    Args:
        query: Query di ricerca
        top_k: Numero di risultati da restituire

    Returns:
        Dict con i risultati della ricerca
    """
    try:
        from llama_index.core import StorageContext, load_index_from_storage

        # Carica l'indice
        storage_context = StorageContext.from_defaults(persist_dir="db_memoria")
        index = load_index_from_storage(storage_context)

        # Crea query engine
        query_engine = index.as_query_engine(similarity_top_k=top_k)

        # Esegue la ricerca
        response = query_engine.query(query)

        # Estrae i nodi sorgente
        source_nodes = []
        if hasattr(response, 'source_nodes'):
            for node in response.source_nodes:
                source_nodes.append({
                    'file_name': node.metadata.get('file_name', 'Unknown'),
                    'score': node.get_score(),
                    'content': node.get_content()[:200] + '...' if len(node.get_content()) > 200 else node.get_content()
                })

        return {
            'status': 'success',
            'query': query,
            'response': str(response),
            'source_nodes': source_nodes,
            'timestamp': datetime.now().isoformat()
        }

    except Exception as e:
        logger.error(f"Errore nella ricerca documenti per query '{query}': {str(e)}")
        return {
            'status': 'error',
            'query': query,
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }

def cleanup_storage(days_old: int = 30) -> Dict[str, Any]:
    """
    Pulisce i file di storage più vecchi di N giorni

    Args:
        days_old: Numero di giorni dopo cui considerare i file vecchi

    Returns:
        Dict con i risultati della pulizia
    """
    try:
        import shutil
        from datetime import timedelta

        cutoff_date = datetime.now() - timedelta(days=days_old)
        cleaned_files = []
        total_size = 0

        # Pulisce i file nella directory db_memoria
        db_memoria_path = "db_memoria"
        if os.path.exists(db_memoria_path):
            for root, dirs, files in os.walk(db_memoria_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    file_modified = datetime.fromtimestamp(os.path.getmtime(file_path))

                    if file_modified < cutoff_date:
                        file_size = os.path.getsize(file_path)
                        total_size += file_size
                        os.remove(file_path)
                        cleaned_files.append(file_path)

        return {
            'status': 'success',
            'cleaned_files': len(cleaned_files),
            'total_size_cleaned': total_size,
            'cutoff_date': cutoff_date.isoformat(),
            'timestamp': datetime.now().isoformat()
        }

    except Exception as e:
        logger.error(f"Errore nella pulizia storage: {str(e)}")
        return {
            'status': 'error',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }
